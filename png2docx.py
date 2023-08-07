import glob
import os
import shutil

from docx import Document
from docx.shared import Cm


def image_2_docx():
    files_name = []
    image_path = os.path.abspath('images_from_pdf_demo')
    os.chdir(image_path)
    for file in glob.glob("*.png"):
        files_name.append(file)
    files_name.sort()

    document = Document()

    for i in range(0, len(files_name), 2):
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(f'{image_path}/{files_name[i]}',
                      width=Cm(15), height=Cm(17))
    os.chdir('..')
    save_path = os.path.abspath('docx_demo')
    document.save(f'{save_path}/demo.docx')

    folder = os.path.abspath('images_from_pdf_demo')  # comment this if you want to see the png files after the convert
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print('Failed to delete %s. Reason: %s' % (file_path, e))
